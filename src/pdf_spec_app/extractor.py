from __future__ import annotations

import importlib
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument

from .extractor_ocr import OcrBackendMixin
from .extractor_pdfplumber import PdfPlumberBackendMixin
from .extractor_pdfplumber import _dedupe_pdfplumber_chars_prefer_latest
from .extractor_power_query import PowerQueryBackendMixin
from .extractor_spire import SpireBackendMixin
from .extractor_support import (
    PROJECT_RUNTIME_DIR,
    ProgressCallback,
    _append_raw_import_text,
    _configure_runtime_environment,
    _describe_pdfplumber_revision_usage,
    _form_markdown_table,
    _group_manual_table_regions_by_page,
    _infer_tables_from_text,
    _split_evaluation_warnings,
    _table_from_html,
    ExtractionError,
)
from .models import ExtractedSegment, ExtractedTable, ExtractionOptions, SourceDocument


_configure_runtime_environment()

try:
    spire_pdf = importlib.import_module("spire.pdf")
except ImportError:  # pragma: no cover - dependency may not be installed in the workspace
    spire_pdf = None

try:
    spire_pdf_common = importlib.import_module("spire.pdf.common")
except ImportError:  # pragma: no cover - dependency may not be installed in the workspace
    spire_pdf_common = None

try:
    import pypdfium2
except ImportError:  # pragma: no cover - dependency may not be installed in the workspace
    pypdfium2 = None

try:
    import pytesseract
except ImportError:  # pragma: no cover - dependency may not be installed in the workspace
    pytesseract = None

try:
    import pdfplumber
except ImportError:  # pragma: no cover - dependency may not be installed in the workspace
    pdfplumber = None


class DocumentExtractor(PdfPlumberBackendMixin, OcrBackendMixin, PowerQueryBackendMixin, SpireBackendMixin):
    """Extracts local document content using multiple local backends."""

    def extract(
        self,
        file_path: str | Path,
        options: ExtractionOptions | None = None,
        progress_callback: ProgressCallback | None = None,
    ) -> SourceDocument:
        self._report(progress_callback, 5.0, "Preparing extraction...")
        options = options or ExtractionOptions()
        path = Path(file_path)
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._extract_pdf(path, options, progress_callback)
        if suffix == ".docx":
            self._report(progress_callback, 15.0, "Reading DOCX document...")
            return self._extract_docx(path)
        if suffix == ".doc":
            raise ExtractionError("Legacy .doc extraction is not supported in the current local-only build.")
        self._report(progress_callback, 25.0, "Reading plain text document...")
        return self._extract_text(path)

    def _extract_docx(self, path: Path) -> SourceDocument:
        try:
            docx_document = DocxDocument(path)
        except Exception as exc:  # pragma: no cover - dependency-specific failures
            raise ExtractionError(f"Failed to read DOCX document: {path.name}") from exc

        paragraphs = [paragraph.text.strip() for paragraph in docx_document.paragraphs if paragraph.text.strip()]
        tables: list[ExtractedTable] = []
        for table in docx_document.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            rows = [row for row in rows if any(cell for cell in row)]
            if not rows:
                continue
            tables.append(
                ExtractedTable(
                    page_number=1,
                    headers=rows[0],
                    rows=rows[1:] if len(rows) > 1 else [],
                    import_headers=list(rows[0]),
                    source_text=_form_markdown_table(rows),
                    raw_text=_form_markdown_table(rows),
                    confidence=0.85,
                )
            )

        raw_docx_text_parts = [paragraph.text for paragraph in docx_document.paragraphs if paragraph.text]
        for table in docx_document.tables:
            for row in table.rows:
                row_values = [cell.text for cell in row.cells]
                raw_docx_text_parts.append("\t".join(row_values))

        document = SourceDocument(
            path=path,
            title=path.stem,
            raw_import_text="\n".join(raw_docx_text_parts),
            tables=tables,
        )
        if paragraphs:
            document.segments.append(
                ExtractedSegment(
                    page_number=1,
                    text="\n\n".join(paragraphs),
                    confidence=0.95,
                    segment_type="docx_text",
                )
            )
        if not document.segments and not document.tables:
            document.segments.append(
                ExtractedSegment(
                    page_number=1,
                    text="[No readable DOCX content was extracted.]",
                    confidence=0.0,
                    segment_type="extraction_gap",
                )
            )
        return document

    def _extract_text(self, path: Path) -> SourceDocument:
        text = path.read_text(encoding="utf-8")
        return SourceDocument(
            path=path,
            title=path.stem,
            raw_import_text=text,
            segments=[ExtractedSegment(page_number=1, text=text, confidence=1.0, segment_type="plain_text")],
        )

    @staticmethod
    def _report(
        progress_callback: ProgressCallback | None,
        percent: float,
        message: str,
    ) -> None:
        if progress_callback is not None:
            progress_callback(percent, message)

    @staticmethod
    def _spire_attr(name: str) -> Any:
        for module in (spire_pdf, spire_pdf_common):
            if module is not None and hasattr(module, name):
                return getattr(module, name)
        raise ExtractionError(f"Spire.PDF does not expose the expected API member: {name}.")

    @staticmethod
    def _spire_pdf_module() -> Any:
        return spire_pdf

    @staticmethod
    def _pdfplumber_module() -> Any:
        return pdfplumber

    @staticmethod
    def _pypdfium2_module() -> Any:
        return pypdfium2

    @staticmethod
    def _pytesseract_module() -> Any:
        return pytesseract
