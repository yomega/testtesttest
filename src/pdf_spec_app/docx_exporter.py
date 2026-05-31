from __future__ import annotations

from pathlib import Path

from .models import Specification

try:
    from docx import Document
except ImportError:  # pragma: no cover - dependency may not be installed in the workspace
    Document = None


class ExportError(RuntimeError):
    """Raised when the Word export cannot be completed."""


class DocxExporter:
    def export(self, specification: Specification, destination: str | Path) -> Path:
        if Document is None:
            raise ExportError("python-docx is not installed. Install project dependencies first.")

        output_path = Path(destination)
        doc = Document()
        doc.add_heading(specification.title, level=1)
        doc.add_paragraph(f"Source document: {specification.source_path}")

        for section in specification.sections:
            doc.add_heading(section.title, level=2)
            for statement in section.statements:
                paragraph = doc.add_paragraph(style="List Bullet")
                paragraph.add_run(statement.text)
                if statement.status != "supported":
                    paragraph.add_run(f" [{statement.status.upper()}]")
                for evidence in statement.evidence:
                    doc.add_paragraph(
                        (
                            f"Page {evidence.page_number} | {evidence.source_type} | "
                            f"confidence {evidence.confidence:.2f} | {evidence.excerpt}"
                        ),
                        style="Intense Quote",
                    )

        doc.save(str(output_path))
        return output_path
